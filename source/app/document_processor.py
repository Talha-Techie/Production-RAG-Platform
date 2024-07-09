"""Document processing service for multiple formats."""
import io
import re
from typing import List, Tuple
from pathlib import Path
import logging

from pypdf import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

from app.config import settings
from app.models import DocumentFormat

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process different document formats and extract text."""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
    
    async def process_document(
        self,
        file_content: bytes,
        file_name: str,
        doc_format: str
    ) -> str:
        """
        Process document and extract text content.
        
        Args:
            file_content: Raw file bytes
            file_name: Original file name
            doc_format: Document format (pdf, txt, md, docx, html)
            
        Returns:
            Extracted text content
        """
        try:
            if doc_format == DocumentFormat.PDF.value:
                return await self._process_pdf(file_content)
            elif doc_format == DocumentFormat.TXT.value:
                return await self._process_txt(file_content)
            elif doc_format == DocumentFormat.MD.value:
                return await self._process_markdown(file_content)
            elif doc_format == DocumentFormat.DOCX.value:
                return await self._process_docx(file_content)
            elif doc_format == DocumentFormat.HTML.value:
                return await self._process_html(file_content)
            else:
                raise ValueError(f"Unsupported format: {doc_format}")
        except Exception as e:
            logger.error(f"Error processing document {file_name}: {e}")
            raise
    
    async def _process_pdf(self, content: bytes) -> str:
        """Extract text from PDF."""
        pdf_file = io.BytesIO(content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n\n"
        
        return text.strip()
    
    async def _process_txt(self, content: bytes) -> str:
        """Extract text from plain text file."""
        try:
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 if utf-8 fails
            return content.decode('latin-1')
    
    async def _process_markdown(self, content: bytes) -> str:
        """Extract text from Markdown file."""
        md_text = content.decode('utf-8')
        # Convert markdown to HTML then extract text
        html = markdown.markdown(md_text)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    async def _process_docx(self, content: bytes) -> str:
        """Extract text from Word document."""
        doc_file = io.BytesIO(content)
        doc = DocxDocument(doc_file)
        
        text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n\n".join(text)
    
    async def _process_html(self, content: bytes) -> str:
        """Extract text from HTML."""
        html = content.decode('utf-8')
        soup = BeautifulSoup(html, 'lxml')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _split_oversized_sentence(self, sentence: str) -> List[str]:
        """Split a sentence longer than chunk_size at word boundaries."""
        parts = []
        while len(sentence) > self.chunk_size:
            split_at = sentence.rfind(' ', 0, self.chunk_size)
            if split_at == -1:
                split_at = self.chunk_size
            parts.append(sentence[:split_at].strip())
            sentence = sentence[split_at:].strip()
        if sentence:
            parts.append(sentence)
        return parts

    def chunk_text(self, text: str, document_name: str = None, document_metadata: dict = None) -> List[Tuple[str, int]]:
        """
        Split text into chunks with overlap.

        Args:
            text: Input text to chunk
            document_name: Optional document name/title to include in chunks
            document_metadata: Optional metadata dict (author, etc.)

        Returns:
            List of (chunk_text, chunk_index) tuples
        """
        # Clean text
        text = self._clean_text(text)

        # Split into sentences (simple approach)
        raw_sentences = re.split(r'[.!?]+\s+', text)

        # Expand any sentence that exceeds chunk_size into word-boundary sub-pieces
        sentences = []
        for s in raw_sentences:
            if len(s) > self.chunk_size:
                sentences.extend(self._split_oversized_sentence(s))
            else:
                sentences.append(s)

        chunks = []
        current_chunk = []
        current_length = 0
        chunk_index = 0

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            sentence_length = len(sentence)

            # If adding this sentence exceeds chunk_size, save current chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                # Enrich chunk with metadata for better searchability
                enriched_chunk = self._enrich_chunk_with_metadata(chunk_text, document_name, document_metadata, chunk_index)
                chunks.append((enriched_chunk, chunk_index))
                chunk_index += 1

                # Keep overlap sentences
                overlap_length = 0
                overlap_sentences = []
                for s in reversed(current_chunk):
                    overlap_length += len(s)
                    if overlap_length > self.chunk_overlap:
                        break
                    overlap_sentences.insert(0, s)

                current_chunk = overlap_sentences
                current_length = overlap_length

            current_chunk.append(sentence)
            current_length += sentence_length

        # Add remaining chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            enriched_chunk = self._enrich_chunk_with_metadata(chunk_text, document_name, document_metadata, chunk_index)
            chunks.append((enriched_chunk, chunk_index))

        return chunks

    def _enrich_chunk_with_metadata(self, chunk_text: str, document_name: str = None,
                                     document_metadata: dict = None, chunk_index: int = 0) -> str:
        """
        Enrich chunk with document metadata for better searchability.

        This adds a header to each chunk with document context, making
        searches for author names, titles, etc. work much better.

        Args:
            chunk_text: Original chunk text
            document_name: Document name/title
            document_metadata: Optional metadata (author, etc.)
            chunk_index: Chunk index for reference

        Returns:
            Enriched chunk text with metadata header
        """
        metadata_parts = []

        # Add document name/title
        if document_name:
            # Clean the document name (remove file extension, etc.)
            clean_name = document_name.replace('.pdf', '').replace('.txt', '')
            metadata_parts.append(f"Document: {clean_name}")

        # Add author if available
        if document_metadata and 'author' in document_metadata:
            metadata_parts.append(f"Author: {document_metadata['author']}")

        # Extract author from filename if it contains "by X" pattern
        if document_name and 'by' in document_name.lower():
            # Try to extract author from patterns like "by Shaw Talebi"
            import re
            match = re.search(r'by\s+([A-Z][a-z]+\s+[A-Z][a-z]+)', document_name)
            if match:
                author = match.group(1)
                metadata_parts.append(f"Author: {author}")

        # If we have metadata, add it as a header
        if metadata_parts:
            header = " | ".join(metadata_parts)
            return f"[{header}]\n\n{chunk_text}"

        return chunk_text
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        # Fix PDF extraction issues with spaces between letters
        # Pattern: "W o r d" → "Word"
        # If there are many single-letter spaces, collapse them
        if text and len(text) > 20:
            # Count non-space characters vs spaces
            char_count = len([c for c in text if not c.isspace()])
            space_count = text.count(' ')

            # If spaces are more than 30% of text, likely has spacing issues
            if space_count > 0 and char_count / (space_count + char_count) < 0.7:
                # Remove all whitespace and rejoin words
                words = ''.join(text.split())
                # Now add back spaces between actual words (capitals indicate word boundaries)
                # Insert space before capital letters (for CamelCase)
                text = re.sub(r'([a-z])([A-Z])', r'\1 \2', words)
                text = re.sub(r'([0-9])([A-Z])', r'\1 \2', text)

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\'"/@]', '', text)
        return text.strip()


# Singleton instance
document_processor = DocumentProcessor()
